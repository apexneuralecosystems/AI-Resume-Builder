"""Text extraction utilities for supported resume/JD file types."""

from pathlib import Path

from fastapi import HTTPException


def extract_text_pdf(file_path: str) -> str:
    """Extract PDF text with resilient multi-pass strategy."""
    import pdfplumber
    chunks: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # First pass: layout-aware extraction (better for multi-column CVs).
            page_text = page.extract_text(
                layout=True,
                x_tolerance=2,
                y_tolerance=2,
            )
            # Fallback pass: plain extraction can recover text missed by layout mode.
            if not page_text or not page_text.strip():
                page_text = page.extract_text()
            # Table pass: append table rows explicitly to avoid missed company/project rows.
            table_lines: list[str] = []
            for table in page.extract_tables() or []:
                for row in table or []:
                    cells = [str(cell).strip() for cell in (row or []) if cell and str(cell).strip()]
                    if cells:
                        table_lines.append(" | ".join(cells))
            if table_lines:
                table_blob = "\n".join(table_lines).strip()
                if page_text and table_blob:
                    page_text = f"{page_text}\n{table_blob}"
                elif table_blob:
                    page_text = table_blob
            # Character fallback: some PDFs have chars but extract_text returns empty.
            if (not page_text or not page_text.strip()) and page.chars:
                line_buckets: dict[int, list[tuple[float, str]]] = {}
                for ch in page.chars:
                    text_val = str(ch.get("text", "") or "").strip()
                    if not text_val:
                        continue
                    top = float(ch.get("top", 0.0) or 0.0)
                    x0 = float(ch.get("x0", 0.0) or 0.0)
                    key = int(round(top / 3.0))  # small tolerance for same visual line
                    line_buckets.setdefault(key, []).append((x0, text_val))
                if line_buckets:
                    lines: list[str] = []
                    for _, chars_in_line in sorted(line_buckets.items(), key=lambda item: item[0]):
                        chars_in_line.sort(key=lambda item: item[0])
                        lines.append("".join(ch for _, ch in chars_in_line).strip())
                    page_text = "\n".join(line for line in lines if line).strip()
            if page_text:
                chunks.append(page_text.strip())
    extracted = "\n\n--- page break ---\n\n".join(chunks).strip()
    if extracted:
        return extracted

    # Final fallback 1: whole-document extraction via pypdf.
    # Helps when page-level extraction APIs return empty/partial content.
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pypdf_pages: list[str] = []
        for page in reader.pages:
            t = (page.extract_text() or "").strip()
            if t:
                pypdf_pages.append(t)
        pypdf_text = "\n\n--- page break ---\n\n".join(pypdf_pages).strip()
        if pypdf_text:
            return pypdf_text
    except Exception:
        # Continue to pdfminer fallback.
        pass

    # Final fallback 2: page text extraction via pypdfium2 textpage API.
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(file_path)
        pdfium_pages: list[str] = []
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            t = (textpage.get_text_range() or "").strip()
            textpage.close()
            page.close()
            if t:
                pdfium_pages.append(t)
        pdf.close()
        pdfium_text = "\n\n--- page break ---\n\n".join(pdfium_pages).strip()
        if pdfium_text:
            return pdfium_text
    except Exception:
        # Continue to pdfminer fallback.
        pass

    # Final fallback 3: whole-document extraction via pdfminer high-level API.
    from pdfminer.high_level import extract_text as pdfminer_extract_text

    return (pdfminer_extract_text(file_path) or "").strip()


def _docx_unique_row_cells_text(row: object) -> list[str]:
    """Row cell texts in order without duplicate merges (common in merged Word tables)."""
    seen_tc: set[int] = set()
    texts: list[str] = []
    for cell in row.cells:
        tc = cell._tc
        if tc in seen_tc:
            continue
        seen_tc.add(tc)
        cleaned = cell.text.strip().replace("\n", " ")
        if cleaned:
            texts.append(cleaned)
    return texts


def extract_text_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t:
            parts.append(t)

    for table_idx, table in enumerate(doc.tables, start=1):
        parts.append("")
        parts.append(f"[Table {table_idx}]")
        for row in table.rows:
            row_parts = _docx_unique_row_cells_text(row)
            if row_parts:
                parts.append(" | ".join(row_parts))
        parts.append("")
    return "\n".join(parts).strip()


def extract_text_doc(file_path: str) -> str:
    """Extract legacy .doc text with fallbacks."""
    import mammoth
    with open(file_path, "rb") as source_file:
        result = mammoth.extract_raw_text(source_file)
    text = (result.value or "").strip()
    if text:
        return text
    # Fallback: decode best-effort bytes when mammoth cannot parse a legacy binary doc.
    with open(file_path, "rb") as source_file:
        raw = source_file.read()
    return raw.decode("utf-8", errors="ignore").strip()


def extract_text(file_path: str, content_type: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    try:
        text = ""
        if ext == ".pdf" or "pdf" in content_type:
            text = extract_text_pdf(file_path)
        elif ext == ".docx" or "openxmlformats" in content_type:
            text = extract_text_docx(file_path)
        elif ext == ".doc" or "msword" in content_type:
            text = extract_text_doc(file_path)
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as source_file:
                text = source_file.read()

        if not text.strip():
            raise HTTPException(
                status_code=422,
                detail=(
                    "Could not extract any text from the file. "
                    "If this is a scanned/image PDF, please upload a text-based PDF/DOCX "
                    "or run OCR on the file first."
                ),
            )
        return text
    except Exception as exc:  # noqa: BLE001
        if isinstance(exc, HTTPException):
            raise
        raise HTTPException(status_code=422, detail=f"Could not extract text from file: {str(exc)}") from exc
